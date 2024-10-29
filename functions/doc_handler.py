from docx import Document
from datetime import datetime
import os
from typing import Dict, List, Optional, Any, Union
from dataclasses import dataclass, field

@dataclass
class ContentNode:
    """Represents a node in the document structure"""
    content: str
    priority: Optional[int] = None
    children: Dict[str, 'ContentNode'] = field(default_factory=dict)
    style_name: str = ''  # Store the original style name for reconstruction

class DocxProcessor:
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.document = Document(file_path)
        # Map priority levels to style names for reconstruction
        self.style_priority_map = self._build_style_priority_map()
        
    def _build_style_priority_map(self) -> Dict[int, str]:
        """Build a mapping of priority levels to style names"""
        style_map = {}
        for style in self.document.styles:
            if hasattr(style, 'priority') and style.priority is not None:
                style_map[style.priority] = style.name
        return style_map

    def _get_style_by_priority(self, priority: int) -> Optional[str]:
        """Get style name for a given priority level"""
        return self.style_priority_map.get(priority)

    def get_structure(self) -> ContentNode:
        """Extract document structure into a ContentNode tree"""
        root = ContentNode(content="ROOT", priority=float('inf'))
        current_path = []  # Stack to track current position in hierarchy
        
        for para in self.document.paragraphs:
            priority = para.style.priority if hasattr(para.style, 'priority') else None
            
            if priority is not None:
                # Pop from current_path until we find a higher priority
                while current_path and current_path[-1].priority <= priority:
                    current_path.pop()
                
                # Create new node
                new_node = ContentNode(
                    content=para.text,
                    priority=priority,
                    style_name=para.style.name
                )
                
                # Add to hierarchy
                parent = current_path[-1] if current_path else root
                parent.children[para.text] = new_node
                current_path.append(new_node)
            else:
                # Regular text - add to most recent section
                if current_path:
                    current_node = current_path[-1]
                    # Use text as key with timestamp to handle multiple paragraphs
                    key = f"{para.text[:30]}_{datetime.now().timestamp()}"
                    current_node.children[key] = ContentNode(
                        content=para.text,
                        priority=None,
                        style_name='Normal'
                    )

        return root

    def update_content(self, section_key: str, new_content: Union[str, Dict[str, Any]], root: Optional[ContentNode] = None) -> bool:
        """
        Update specific sections of the document while preserving others.
        Returns True if update was successful.
        
        Parameters:
            section_key: Key of section to update
            new_content: New content (either string for simple update or dict for structural update)
            root: Root node (optional, used for recursion)
        """
        if root is None:
            root = self.get_structure()

        def update_node_recursive(node: ContentNode, content_dict: Dict[str, Any]) -> ContentNode:
            """Recursively update node with new content structure"""
            updated_node = ContentNode(
                content=content_dict.get('content', node.content),
                priority=content_dict.get('priority', node.priority),
                style_name=content_dict.get('style_name', node.style_name)
            )
            
            # Update children if provided
            if 'children' in content_dict:
                for child_key, child_content in content_dict['children'].items():
                    if isinstance(child_content, dict):
                        updated_node.children[child_key] = update_node_recursive(
                            ContentNode("", None), child_content
                        )
                    else:
                        updated_node.children[child_key] = ContentNode(
                            content=child_content,
                            priority=None,
                            style_name='Normal'
                        )
            
            return updated_node

        def find_and_update(current_node: ContentNode, target_key: str, new_data: Union[str, Dict[str, Any]]) -> bool:
            if target_key in current_node.children:
                if isinstance(new_data, str):
                    current_node.children[target_key].content = new_data
                else:
                    current_node.children[target_key] = update_node_recursive(
                        current_node.children[target_key], new_data
                    )
                return True
                
            for child in current_node.children.values():
                if find_and_update(child, target_key, new_data):
                    return True
            return False

        return find_and_update(root, section_key, new_content)

    def _rebuild_document(self, root: ContentNode) -> None:
        """Rebuild document from ContentNode structure"""
        self.document = Document()
        
        def add_node_to_document(node: ContentNode) -> None:
            if node.content != "ROOT":
                paragraph = self.document.add_paragraph(node.content)
                if node.priority is not None:
                    paragraph.style = node.style_name
                
            # Add all children in order of priority
            sorted_children = sorted(
                node.children.values(),
                key=lambda x: (x.priority or -1, x.content),
                reverse=True
            )
            
            for child in sorted_children:
                add_node_to_document(child)
                
        add_node_to_document(root)

    def save_document(self, root: Optional[ContentNode] = None) -> str:
        """Save document with date in filename"""
        if root:
            self._rebuild_document(root)
            
        file_dir = os.path.dirname(self.file_path)
        base_name = os.path.splitext(os.path.basename(self.file_path))[0]
        date_str = datetime.now().strftime('%Y%m%d')
        new_filename = f"{base_name}_{date_str}.docx"
        new_path = os.path.join(file_dir, new_filename)
        
        self.document.save(new_path)
        return new_path

# Example usage:
if __name__ == "__main__":
    processor = DocxProcessor("example.docx")
    
    # Get current structure
    doc_structure = processor.get_structure()
    
    # Example of updating a specific section
    new_section_content = {
        "content": "Updated Section Title",
        "priority": 9,
        "style_name": "Heading 1",
        "children": {
            "subsection1": {
                "content": "New Subsection",
                "priority": 8,
                "style_name": "Heading 2",
                "children": {
                    "text1": "New content paragraph"
                }
            }
        }
    }
    
    processor.update_content("Original Section Title", new_section_content, doc_structure)
    processor.save_document(doc_structure)